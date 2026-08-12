"""build the paper as an IEEE Computer Society two-column .docx.

full-width title block, then a continuous section break into two columns for the
body. Times New Roman, 10pt body, numbered [n] references, figures at column
width, ablation table, clickable repository link.
"""
import os

import docx
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
FONT = "Times New Roman"
OUT = os.path.join(HERE, "paper.docx")

REPO = "github.com/sunjungco2027-hub/cgm-individuality"
REPO_URL = "https://github.com/sunjungco2027-hub/cgm-individuality"


def set_columns(section, num, space_dxa):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_dxa))


def margins(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)


def run(p, text, size=10, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def para(container, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
         indent=None, line=None):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    if line is not None:
        pf.line_spacing = line
    return p


def add_hyperlink(p, text, url, size=10):
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), FONT)
    rf.set(qn("w:hAnsi"), FONT)
    rPr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    link.append(r)
    p._p.append(link)


def body_par(doc, text, do_indent):
    p = para(doc, indent=0.2 if do_indent else None, after=0)
    run(p, text, 10)
    return p


def heading(doc, num, title, italic=False):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=2)
    p.paragraph_format.keep_with_next = True
    run(p, (num + " " if num else "") + title, 10, bold=True, italic=italic)


def figure(doc, fname, caption):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
    p.add_run().add_picture(os.path.join(FIG, fname), width=Inches(3.3))
    c = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, after=6)
    run(c, caption, 8)


DOC = Document()

# ----- section 1: full-width title block (single column) -----
s1 = DOC.sections[0]
margins(s1)
set_columns(s1, 1, 288)

title = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6)
run(title, "Covariate-Independent Individuality and Subject Re-Identification "
           "in CGM Postprandial Glucose Excursions", 20)
auth = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
run(auth, "Sun Jung", 11)
aff = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
run(aff, "International School Ho Chi Minh City", 10, italic=True)
role = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
run(role, "High School Student, Class of 2027", 10, italic=True)
em = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
run(em, "sun.jung.co2027@gmail.com", 10)

# ----- section 2: two-column body -----
s2 = DOC.add_section(WD_SECTION.CONTINUOUS)
margins(s2)
set_columns(s2, 2, 288)

ABSTRACT = ("Continuous glucose monitors record how a person responds to every "
    "meal, but it is unclear how much of that response is specific to the person "
    "rather than to the food. Using the public CGMacros dataset, we describe each "
    "of 1,657 meal excursions from 45 subjects with ten features covering "
    "amplitude, timing, and shape, and treat each meal as a repeated measurement "
    "of its subject. After adjusting for demographics and fasting labs, peak and "
    "baseline glucose remain the most person-specific features, meals from the "
    "same person cluster together more tightly than meals from different people "
    "(permutation p = 0.003), and a nearest-neighbor model re-identifies the "
    "subject behind a held-out meal well above chance (the correct person is among "
    "the top five guesses 39% of the time, against 11% by chance). The same "
    "features predict diabetes under subject-grouped cross-validation, and the "
    "features that are most individual are also the most predictive. Postprandial "
    "glucose responses therefore carry a covariate-independent individual "
    "signature that doubles as a signal of metabolic health.")
ab = para(DOC, after=4)
run(ab, "Abstract—", 9, bold=True, italic=True)
run(ab, ABSTRACT, 9, bold=True)

it = para(DOC, after=8)
run(it, "Index Terms—", 9, bold=True, italic=True)
run(it, "continuous glucose monitoring, individuality, re-identification, "
        "intraclass correlation, diabetes classification", 9, bold=True)

# blocks: ('h1'/'h2', num, title) | ('p', text) | ('fig', file, cap) | ('table',)
BLOCKS = [
    ("h1", "1.", "Introduction"),
    ("p", "Continuous glucose monitors are now cheap enough that people without "
     "diabetes wear them, and each sensor produces a dense record of how the body "
     "handles food. The most informative part of that record is the meal "
     "response: glucose rises after eating, reaches a peak, and returns toward "
     "baseline. The height, timing, and steepness of that curve differ from person "
     "to person, and this paper asks a simple question about those differences. Is "
     "the shape of a meal response individual to the person, in a way that ordinary "
     "explanations like age and body weight cannot account for, and if so, does "
     "that same individual signal say anything about their metabolic health?"),
    ("p", "Earlier work has shown that people mount different responses to the "
     "same food and that continuous glucose patterns fall into distinct groups. "
     "What it has largely left untested is whether the individual component "
     "survives once demographics and laboratory values are removed, and it has not "
     "tried to measure that individuality the way biometrics does, by asking "
     "whether an unseen meal can be traced back to the right person. Framing the "
     "problem as re-identification is what lets us report individuality as an "
     "accuracy rather than an impression. Cast this way it is a data-mining "
     "question: finding a stable individual signature in noisy, repeated "
     "physiological measurements, and mining that signature for a downstream "
     "clinical signal."),
    ("p", "This paper makes two contributions. First, using covariate-adjusted "
     "repeatability, multivariate distance, and a re-identification test on "
     "held-out meals, we show that postprandial excursion profiles carry an "
     "individual signal that demographics and labs do not explain. Second, we show "
     "that the same features that make a person identifiable also predict diabetes "
     "under subject-grouped cross-validation, so individuality and clinical signal "
     "turn out to live in the same features rather than in separate ones."),

    ("h1", "2.", "Related Work"),
    ("p", "Continuous glucose monitoring records interstitial glucose every few "
     "minutes and has changed how day-to-day metabolism is observed [1]. In "
     "routine use, though, most of that signal is compressed into a handful of "
     "static summaries. Time in range and the glucose management indicator report "
     "how often glucose sits inside a target band or estimate an A1c-like average "
     "[2], [3], and variability indices add a measure of spread [4]. These numbers "
     "are useful for management, but they mostly describe where glucose is rather "
     "than how it moves. A meal response has a shape: a rise, a peak, and a return "
     "toward baseline, and that shape is largely lost once a day is reduced to an "
     "average. Clinical guidelines have argued for years that the postprandial "
     "period itself carries information about metabolic health [5], which suggests "
     "the excursion is worth modeling on its own terms."),
    ("p", "That the same meal can produce very different responses in different "
     "people is by now well documented. Zeevi et al. [6] found that postprandial "
     "responses to identical foods vary widely across individuals and can be "
     "predicted from personal and microbiome features, which motivated "
     "personalized rather than uniform dietary advice. Hall et al. [7] took the "
     "idea further, clustering CGM traces into distinct “glucotypes” that "
     "separated people by how variable their glucose was, including people without "
     "a diabetes diagnosis. Both results establish that individual differences are "
     "real. What they leave open is whether those differences are stable and "
     "specific enough to act as a signature of the person once the obvious "
     "explanations, such as age, body mass, and blood chemistry, are accounted "
     "for."),
    ("p", "Measuring how much of a trait belongs to the person is an old problem. "
     "The intraclass correlation coefficient (ICC) splits the variance of repeated "
     "measurements into a between-subject part and a within-subject part, so a high "
     "ICC means a feature is consistent within a person yet differs across people. "
     "A related but stricter view comes from biometrics, where the question is not "
     "only whether people differ on average but whether an unseen sample can be "
     "matched back to the right individual. Posing CGM meal responses this way, as "
     "a re-identification task, turns a loose notion of individuality into a "
     "measurable quantity: how often the correct person appears among the top "
     "guesses for a held-out meal. To our knowledge this framing has not been "
     "applied to postprandial glucose curves, even though recognizing a person "
     "from a physiological signal is a familiar goal elsewhere."),
    ("p", "CGM-derived features have also been used to separate diabetic from "
     "non-diabetic individuals, which is consistent with the underlying "
     "physiology: impaired insulin secretion and insulin resistance push peaks "
     "higher and slow the return to baseline after eating [8]. One caution runs "
     "through this line of work. When the diabetes label is itself defined by "
     "fasting glucose or A1c, features that essentially re-measure glucose level "
     "will track that label almost by construction. The more informative question "
     "is whether the shape and timing of the excursion, and not only its height, "
     "add anything beyond the diagnostic thresholds. We keep this distinction in "
     "mind when we later separate amplitude features from shape and kinetic ones."),

    ("h1", "3.", "Methods"),
    ("h2", "3.1.", "Data and feature extraction"),
    ("p", "We use the public CGMacros dataset, which pairs minute-level CGM traces "
     "with meal logs and a per-subject panel of demographics and fasting labs. A "
     "meal is any logged event with positive caloric intake. Around each meal we "
     "take the window from 30 minutes before to 240 minutes after and index it in "
     "minutes relative to the meal, so every excursion is described on the same "
     "time axis (Figure 1). After dropping windows with too little coverage, this "
     "yields 1,657 meal excursions from 45 subjects."),
    ("fig", "fig-window.png", "Figure 1.  An example postprandial excursion, the "
     "mean meal response of one subject, with the pre-meal baseline and the peak "
     "marked. The ten features are read off this curve."),
    ("p", "From each window we compute ten features that summarize the curve rather "
     "than a single average. Four are amplitude features (pre-meal baseline, peak, "
     "the rise from baseline to peak, and the area above baseline), two are timing "
     "features (time to peak and time back to baseline), two are kinetic (the "
     "average up-slope and down-slope), and two describe the distribution of "
     "glucose in the window (skewness and kurtosis). The result is one row per "
     "meal, which the later analyses treat as a repeated measurement of the person "
     "who ate it."),
    ("h2", "3.2.", "Covariate-adjusted individuality"),
    ("p", "For each feature we first ask how much of its variance is between people "
     "rather than within a person across meals. We use the intraclass correlation "
     "coefficient in its one-way form, computed from the between- and "
     "within-subject mean squares with a correction for the uneven number of meals "
     "per subject. A raw ICC near one means the feature is highly repeatable within "
     "a person and separates people well."),
    ("p", "A raw ICC can partly reflect demographics rather than anything truly "
     "individual, since older or heavier people tend to run higher glucose and that "
     "shows up as between-subject variance. To separate the two, we regress each "
     "feature on a set of covariates and recompute the ICC on the residuals. The "
     "covariates are the fourteen numeric fields in the subject panel (age, "
     "body-mass index, weight, height, HbA1c, fasting glucose, insulin, and the "
     "full lipid profile) together with one-hot encodings of sex and "
     "self-identified ethnicity. What survives this adjustment is individuality "
     "that age, body size, and blood chemistry do not explain. We attach a 95% "
     "confidence interval to each adjusted ICC with a cluster bootstrap that "
     "resamples subjects with replacement 800 times, keeping each person's meals "
     "together."),
    ("h2", "3.3.", "Multivariate fingerprinting"),
    ("p", "The ICC looks at one feature at a time. To ask whether the whole "
     "excursion profile clusters by person, we residualize all ten features "
     "against the same covariates, standardize them, and measure the cosine "
     "distance between every pair of meals. If meals carry an individual "
     "signature, two meals from the same person should sit closer together than "
     "two meals from different people. We compare the mean within-subject distance "
     "to the mean between-subject distance and summarize the gap with Cohen's d. "
     "Because the pairs are not independent, we test the gap with a permutation "
     "test: we shuffle the subject labels 300 times and count how often a "
     "reshuffled gap is at least as large as the observed one."),
    ("h2", "3.4.", "Re-identification"),
    ("p", "Distances tell us that people separate, but not by how much in "
     "practical terms. To put a number on it we set up a re-identification task. "
     "Each subject's meals are split, 80 percent into a gallery and 20 percent "
     "held out. A nearest-neighbor model with ten neighbors and a cosine metric is "
     "fit on the residualized gallery features, and we ask it to name the subject "
     "behind each held-out meal. We report top-1 accuracy, whether the correct "
     "person is the first guess, and top-5 accuracy, whether the correct person is "
     "among the first five, each against the chance rates of 1/45 and 5/45. We "
     "repeat this over 20 random gallery and probe splits and report the mean."),
    ("h2", "3.5.", "Diabetes classification"),
    ("p", "Finally we test whether the same features predict diabetes. Subjects "
     "are labeled diabetic if their HbA1c is at least 6.5 percent or their fasting "
     "glucose is at least 126 mg/dL, following the ADA thresholds, and the label "
     "is carried down to every meal from that subject. A logistic-regression "
     "classifier is trained on the meal-level features after median imputation and "
     "standardization. Because meals from one person are correlated, a random "
     "split would let the model recognize the person instead of the condition, so "
     "we evaluate with five-fold cross-validation grouped by subject, meaning every "
     "subject's meals fall entirely in either the training or the test fold. We "
     "read feature importance from the standardized logistic-regression "
     "coefficients. To test whether the excursion shape predicts beyond glucose "
     "level, we also re-run this classifier on feature subsets that drop the "
     "amplitude features."),
    ("h2", "3.6.", "Implementation"),
    ("p", "All analyses use Python with scikit-learn [9]. Feature extraction, the "
     "ICC and fingerprinting analyses, the re-identification model, and the "
     "classifier are each a small script, and a single command runs the full "
     "pipeline and prints the summary numbers. The CGMacros data is public but is "
     "not redistributed here."),

    ("h1", "4.", "Results"),
    ("h2", "4.1.", "Cohort"),
    ("p", "The pipeline extracts 1,657 meal excursions from 45 subjects. Applying "
     "the ADA thresholds at the subject level labels 18 of the 45 as diabetic, so "
     "about 40 percent of the cohort, and every meal inherits its subject's label. "
     "The classification task below is therefore moderately imbalanced but not "
     "severely so."),
    ("h2", "4.2.", "Individuality of single features"),
    ("p", "The amplitude features are the most person-specific. Peak glucose has "
     "the highest raw ICC at 0.65 and baseline glucose is next at 0.56, meaning "
     "that most of the variance in these features is between people rather than "
     "within a person across meals. Timing and distribution-shape features sit "
     "much lower, mostly under 0.15. Adjusting for the covariates shrinks every "
     "ICC, as expected, but does not erase the leaders: peak and baseline glucose "
     "keep adjusted ICCs of 0.24 (95% CI 0.15 to 0.33) and 0.20 (0.10 to 0.30). In "
     "other words, part of what makes two people's meal responses differ in height "
     "is not explained by their age, body size, or blood chemistry."),
    ("h2", "4.3.", "Profiles cluster by person"),
    ("p", "Looking at all ten features together tells the same story. After "
     "residualizing against the covariates, the mean cosine distance between two "
     "meals from the same person is 0.912, smaller than the 0.998 between meals "
     "from different people. The gap is a small effect by Cohen's conventions, d = "
     "0.18, but the permutation test rules out chance, with p = 0.003 across 300 "
     "shuffles of the subject labels. So the individual signal is modest in size "
     "yet reliably present once demographics and labs are removed (Figure 2)."),
    ("fig", "fig1-distance.png", "Figure 2.  Cosine distances between meals in the "
     "residualized feature space. Within-subject pairs sit closer than "
     "between-subject pairs."),
    ("h2", "4.4.", "Re-identification"),
    ("p", "The re-identification task turns that signal into a concrete number. "
     "Given a held-out meal, the nearest-neighbor model names the correct person "
     "first about 15 percent of the time, against a chance rate of about 2 "
     "percent, and it places the correct person in its top five guesses about 39 "
     "percent of the time, against a chance rate of 11 percent (means over 20 "
     "random splits). Both are several times chance. The accuracy is far from "
     "identification-grade, which is expected for a small set of features, but it "
     "shows that a meal response carries enough of a signature to point back at "
     "the person who produced it."),
    ("h2", "4.5.", "Diabetes classification"),
    ("p", "The same features predict diabetes reasonably well. Under five-fold "
     "subject-grouped cross-validation the logistic-regression classifier reaches "
     "an area under the ROC curve of 0.87, so the features generalize to people "
     "the model has not seen rather than memorizing individuals. The standardized "
     "coefficients rank peak glucose first, then baseline glucose and the rise "
     "from baseline to peak. The amplitude features do most of the work, with "
     "timing and shape features contributing little on their own (Figure 3)."),
    ("fig", "fig2-importance.png", "Figure 3.  Standardized logistic-regression "
     "coefficients. Peak and baseline glucose carry the most weight."),
    ("h2", "4.6.", "The two results meet on the same features"),
    ("p", "Peak and baseline glucose are the two features that keep the most "
     "individuality after covariate adjustment, and they are also the two that "
     "carry the most weight in the diabetes classifier. The trait that makes a "
     "person recognizable is, at least here, the same trait that flags their "
     "metabolic risk. That overlap is worth a caution as much as a claim: because "
     "the diabetes label is defined by glucose level, features that re-measure "
     "level are expected to predict it, so the honest open question is how much the "
     "shape of the excursion adds beyond height. To test this directly, we re-ran "
     "the classifier on feature subsets (Table 1). Removing peak and baseline "
     "glucose lowers the AUC from 0.87 to 0.71, and dropping every amplitude "
     "feature, leaving only timing, kinetic, and shape, still reaches 0.66, well "
     "above the 0.5 chance level. Glucose level alone gives 0.85. Amplitude "
     "therefore carries most of the signal, as the label makes inevitable, but the "
     "shape of the excursion predicts diabetes on its own."),
    ("table",),

    ("h1", "5.", "Discussion"),
    ("p", "Two findings sit next to each other here. The shape of a meal response "
     "holds an individual component that survives adjustment for demographics and "
     "fasting labs, and the features that carry that component are the same ones "
     "that predict diabetes. The most likely reading is that amplitude features "
     "such as peak and baseline glucose reflect slower-moving traits of a person's "
     "metabolism, for instance insulin secretion and clearance, that a one-time "
     "demographic and lab panel does not fully capture. Latent factors of that "
     "kind, and possibly gastric emptying or gut microbiome, would show up as "
     "stable between-person differences in how high glucose climbs and how it "
     "settles, which is what the ICC and the fingerprint analyses pick up. For the "
     "growing number of people who already wear a sensor, this points toward a "
     "passive form of metabolic characterization that needs no extra visit."),
    ("h2", "5.1.", "Limitations"),
    ("p", "The study is small. Forty-five subjects and a single dataset limit how "
     "far the numbers generalize, and the cohort leans toward one population, so "
     "the same analysis on a broader group could look different. The most important "
     "caveat is the label itself. Diabetes here is defined by fasting glucose and "
     "HbA1c, and the two strongest predictors re-measure glucose level, so part of "
     "the classification result is expected almost by definition. The ablation in "
     "Table 1 speaks to this: with the amplitude features removed, the shape and "
     "kinetic features still predict above chance, so the result is not only a "
     "restatement of glucose level. The re-identification accuracy is well above "
     "chance but modest, and it moves a little with the random split, so it should "
     "be read as evidence of a signature rather than as a working identifier. "
     "Finally, the ten features are a compact summary, and meal composition, time "
     "of day, and overlap between nearby meals are not modeled, all of which could "
     "sharpen or blur the individual signal."),

    ("h1", "6.", "Conclusion"),
    ("p", "Postprandial glucose responses carry an individual signature that is "
     "not just a restatement of a person's demographics or labs, and that same "
     "signature aligns with their metabolic risk. The effect is small and the "
     "cohort is limited, but it is consistent across a repeatability measure, a "
     "distance-based test, and a re-identification task, and it survives covariate "
     "adjustment. Larger and more diverse data, richer features, and the "
     "amplitude-free test noted above would show how much further the idea holds."),

    ("h1s", "Reproducibility"),
    ("repro",),
    ("h1s", "Acknowledgment"),
    ("p", "The author thanks Dr. Seo Ho Song for guidance and feedback throughout "
     "this project. This work uses the CGMacros dataset, and the author thanks its "
     "creators and PhysioNet for making the data openly available."),
]

flush_next = False  # IEEE: first paragraph after a heading is not indented
for b in BLOCKS:
    kind = b[0]
    if kind in ("h1", "h2"):
        heading(DOC, b[1], b[2], italic=(kind == "h2"))
        flush_next = True
    elif kind == "h1s":
        heading(DOC, "", b[1])
        flush_next = True
    elif kind == "p":
        body_par(DOC, b[1], do_indent=not flush_next)
        flush_next = False
    elif kind == "fig":
        figure(DOC, b[1], b[2])
        flush_next = False
    elif kind == "repro":
        p = para(DOC, indent=None if flush_next else 0.2)
        run(p, "The CGMacros dataset is publicly available from PhysioNet. All "
               "feature-extraction and analysis code is open source at ", 10)
        add_hyperlink(p, REPO, REPO_URL, 10)
        run(p, ", and a single command regenerates every number reported here. "
               "The pipeline is deterministic: random splits are seeded and the "
               "re-identification result is averaged over twenty fixed splits. A "
               "test suite pins the reported values so that any change moving them "
               "is caught, and a separate set of tests checks the intraclass "
               "correlation against an independent one-way ANOVA computation and "
               "verifies the effect-size and permutation calculations on synthetic "
               "data with known answers. The results in this paper correspond to "
               "the tagged release v1.0 of the repository.", 10)
        flush_next = False
    elif kind == "table":
        cap = para(DOC, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
        run(cap, "TABLE 1.  Diabetes prediction under feature ablation "
                 "(subject-grouped CV).", 8)
        rows = [
            ("Feature set", "AUC"),
            ("All ten features", "0.87"),
            ("Without peak and baseline", "0.71"),
            ("Without any amplitude feature", "0.66"),
            ("Peak and baseline only", "0.85"),
        ]
        tbl = DOC.add_table(rows=len(rows), cols=2)
        tbl.columns[0].width = Inches(2.6)
        tbl.columns[1].width = Inches(0.7)
        for ri, (a, c) in enumerate(rows):
            for ci, val in enumerate((a, c)):
                cell = tbl.cell(ri, ci)
                cell.width = Inches(2.6 if ci == 0 else 0.7)
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_after = Pt(0)
                run(cp, val, 8, bold=(ri == 0))
        # horizontal rules: top of table, under header, bottom
        def hrule(cell, where):
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            e = OxmlElement("w:" + where)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")
            borders.append(e)
            tcPr.append(borders)
        for ci in range(2):
            hrule(tbl.cell(0, ci), "top")
            hrule(tbl.cell(0, ci), "bottom")
            hrule(tbl.cell(len(rows) - 1, ci), "bottom")
        flush_next = False

# ----- references -----
heading(DOC, "", "References")
REFS = [
    "T. Danne, R. Nimri, T. Battelino, et al., “International consensus on use "
    "of continuous glucose monitoring,” Diabetes Care, vol. 40, no. 12, pp. "
    "1631–1640, 2017.",
    "T. Battelino, T. Danne, R. M. Bergenstal, et al., “Clinical targets for "
    "continuous glucose monitoring data interpretation,” Diabetes Care, vol. "
    "42, no. 8, pp. 1593–1603, 2019.",
    "R. M. Bergenstal, R. W. Beck, K. L. Close, et al., “Glucose management "
    "indicator (GMI): a new term for estimating A1C from continuous glucose "
    "monitoring,” Diabetes Care, vol. 41, no. 11, pp. 2275–2280, 2018.",
    "L. Monnier, C. Colette, and D. R. Owens, “Glycemic variability: the third "
    "component of the dysglycemia in diabetes,” J. Diabetes Sci. Technol., "
    "vol. 2, no. 6, pp. 1094–1100, 2008.",
    "A. Ceriello and S. Colagiuri, “International Diabetes Federation "
    "guideline for management of postmeal glucose,” Diabetic Medicine, vol. "
    "25, no. 10, pp. 1151–1156, 2008.",
    "D. Zeevi, T. Korem, N. Zmora, et al., “Personalized nutrition by "
    "prediction of glycemic responses,” Cell, vol. 163, no. 5, pp. "
    "1079–1094, 2015.",
    "H. Hall, D. Perelman, A. Breschi, et al., “Glucotypes reveal new patterns "
    "of glucose dysregulation,” PLoS Biology, vol. 16, no. 7, e2005143, 2018.",
    "R. A. DeFronzo, “Pathogenesis of type 2 diabetes mellitus,” Medical "
    "Clinics of North America, vol. 88, no. 4, pp. 787–835, 2004.",
    "F. Pedregosa, G. Varoquaux, A. Gramfort, et al., “Scikit-learn: machine "
    "learning in Python,” J. Mach. Learn. Res., vol. 12, pp. 2825–2830, "
    "2011.",
]
for i, r in enumerate(REFS, 1):
    p = para(DOC, after=2)
    p.paragraph_format.left_indent = Inches(0.17)
    p.paragraph_format.first_line_indent = Inches(-0.17)
    run(p, "[%d]  %s" % (i, r), 8)

DOC.save(OUT)
print("wrote", OUT)
